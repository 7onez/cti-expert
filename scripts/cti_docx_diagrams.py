"""
CTI Report Diagrams — entity relationship and network topology via networkx + matplotlib.
"""
import matplotlib
matplotlib.use("Agg")

import numpy as np
# NetworkX 2.x uses np.alltrue removed in NumPy 2.0 — patch for compatibility
if not hasattr(np, "alltrue"):
    np.alltrue = np.all

import matplotlib.pyplot as plt
import networkx as nx
from io import BytesIO
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from cti_docx_styles import COLORS_HEX


# Entity type → color mapping
ENTITY_COLORS = {
    "person": "#8C2D2D",        # brick — the anchor a reader finds first
    "username": "#3B5566",      # steel
    "email": "#5A4A7A",         # muted violet
    "domain": "#B0790F",        # ochre
    "ip": "#6F6A61",            # muted
    "organization": "#5A6B3B",  # olive
    "phone": "#7A2F52",         # muted plum
    "location": "#9A5B2F",      # amber-brown
    "asset": "#22333F",         # slate
    "event": "#2F6B6B",         # teal
}

# Connection type → edge style
EDGE_STYLES = {
    "owns": {"style": "solid", "color": "#8C2D2D", "width": 2.5},
    "uses": {"style": "solid", "color": "#3B5566", "width": 1.5},
    "works_at": {"style": "dashed", "color": "#5A6B3B", "width": 1.5},
    "linked_to": {"style": "dotted", "color": "#6F6A61", "width": 1.0},
    "alias": {"style": "dashdot", "color": "#5A4A7A", "width": 1.5},
    "communicated_with": {"style": "solid", "color": "#7A2F52", "width": 1.5},
}

ENTITY_ICONS = {
    "person": "[P]",
    "username": "[@]",
    "email": "[E]",
    "domain": "[D]",
    "ip": "[IP]",
    "organization": "[O]",
    "phone": "[Ph]",
    "location": "[L]",
    "asset": "[A]",
    "event": "[Ev]",
}


def _save_fig_to_buffer(fig) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150, facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf


def _embed_editorial(doc, kind: str, ctx: dict) -> bool:
    """Try the Diagram Design editorial SVG (rasterized via cairosvg). True on success.

    Returns False when the generator or cairosvg/cairo is unavailable, so callers
    fall back to the matplotlib figure. This is what makes the DOCX diagrams match
    the editorial HTML/PDF figures byte-for-byte in layout.
    """
    try:
        from cti_diagram_design import build_entity_svg, build_topology_svg, render_svg_to_png
    except Exception:
        return False
    try:
        svg = build_entity_svg(ctx) if kind == "entity" else build_topology_svg(ctx)
        png = render_svg_to_png(svg, scale=2.0)
        if not png:
            return False
        doc.add_picture(BytesIO(png), width=Inches(6.4))
        doc.paragraphs[-1].alignment = 1
        return True
    except Exception:
        return False


def add_entity_diagram(doc, subjects: list, connections: list, data: dict = None) -> None:
    """Render entity relationship diagram — editorial SVG first, networkx fallback."""
    if not subjects:
        return
    ctx = data if isinstance(data, dict) and data.get("subjects") else {
        "subjects": subjects, "connections": connections}
    if _embed_editorial(doc, "entity", ctx):
        return
    if not connections:
        return

    G = nx.DiGraph()

    # Add nodes
    node_colors = []
    node_labels = {}
    for s in subjects:
        sid = s.get("id", s.get("label", "?"))
        stype = s.get("type", "person").lower()
        label = s.get("label", sid)
        icon = ENTITY_ICONS.get(stype, "?")
        G.add_node(sid)
        node_labels[sid] = f"{icon} {label}"
        node_colors.append(ENTITY_COLORS.get(stype, "#6F6A61"))

    # Add edges
    edge_colors = []
    edge_styles = []
    edge_widths = []
    edge_labels = {}
    for c in connections:
        from_id = c.get("from_id", "")
        to_id = c.get("to_id", "")
        rel = c.get("relationship", "linked_to")
        if from_id in G.nodes and to_id in G.nodes:
            G.add_edge(from_id, to_id)
            style_info = EDGE_STYLES.get(rel, EDGE_STYLES["linked_to"])
            edge_colors.append(style_info["color"])
            edge_styles.append(style_info["style"])
            edge_widths.append(style_info["width"])
            edge_labels[(from_id, to_id)] = rel

    if len(G.nodes) == 0:
        return

    # Layout
    fig, ax = plt.subplots(figsize=(8, 6), dpi=150)
    pos = nx.spring_layout(G, seed=42, k=2.0)

    # Draw edges
    for i, (u, v) in enumerate(G.edges()):
        nx.draw_networkx_edges(
            G, pos, edgelist=[(u, v)], ax=ax,
            edge_color=[edge_colors[i]] if i < len(edge_colors) else ["#6F6A61"],
            width=edge_widths[i] if i < len(edge_widths) else 1.0,
            style=edge_styles[i] if i < len(edge_styles) else "solid",
            arrows=True, arrowsize=15, arrowstyle="-|>",
            connectionstyle="arc3,rad=0.1"
        )

    # Draw nodes
    nx.draw_networkx_nodes(
        G, pos, ax=ax,
        node_color=node_colors,
        node_size=2000,
        alpha=0.9,
        edgecolors="white",
        linewidths=2
    )

    # Draw labels
    nx.draw_networkx_labels(
        G, pos, labels=node_labels, ax=ax,
        font_size=7, font_color="white", font_weight="bold"
    )

    # Draw edge labels
    nx.draw_networkx_edge_labels(
        G, pos, edge_labels=edge_labels, ax=ax,
        font_size=6, font_color=COLORS_HEX["muted"],
        bbox=dict(boxstyle="round,pad=0.2", facecolor="white", alpha=0.8)
    )

    ax.set_title("Entity Relationship Map", fontsize=14, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=20)
    ax.axis("off")

    # Legend
    legend_items = []
    seen_types = set()
    for s in subjects:
        stype = s.get("type", "person").lower()
        if stype not in seen_types:
            seen_types.add(stype)
            icon = ENTITY_ICONS.get(stype, "?")
            color = ENTITY_COLORS.get(stype, "#6F6A61")
            legend_items.append(
                plt.Line2D([0], [0], marker="o", color="w",
                          markerfacecolor=color, markersize=8,
                          label=f"{icon} {stype.title()}")
            )
    if legend_items:
        ax.legend(handles=legend_items, loc="lower left", fontsize=7,
                 framealpha=0.9, edgecolor=COLORS_HEX["border"])

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(6))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


def add_network_topology(doc, subjects: list, connections: list, data: dict = None) -> None:
    """Render network topology — editorial SVG first, networkx fallback."""
    if subjects:
        ctx = data if isinstance(data, dict) and data.get("subjects") else {
            "subjects": subjects, "connections": connections}
        if _embed_editorial(doc, "topology", ctx):
            return
    infra_types = {"domain", "ip", "organization"}
    infra_subjects = [s for s in subjects if s.get("type", "").lower() in infra_types]
    infra_ids = {s.get("id", s.get("label", "")) for s in infra_subjects}

    infra_connections = [
        c for c in connections
        if c.get("from_id") in infra_ids or c.get("to_id") in infra_ids
    ]

    if not infra_subjects:
        # Fallback: use all subjects
        infra_subjects = subjects
        infra_connections = connections

    if len(infra_subjects) < 2:
        return

    G = nx.Graph()
    node_colors = []
    node_labels = {}

    for s in infra_subjects:
        sid = s.get("id", s.get("label", "?"))
        stype = s.get("type", "domain").lower()
        label = s.get("label", sid)
        G.add_node(sid)
        node_labels[sid] = f"{label}"
        node_colors.append(ENTITY_COLORS.get(stype, "#6F6A61"))

    for c in infra_connections:
        from_id = c.get("from_id", "")
        to_id = c.get("to_id", "")
        if from_id in G.nodes and to_id in G.nodes:
            G.add_edge(from_id, to_id, label=c.get("relationship", ""))

    fig, ax = plt.subplots(figsize=(7, 5), dpi=150)
    pos = nx.spring_layout(G, seed=123, k=2.0)

    nx.draw_networkx_edges(G, pos, ax=ax, edge_color=COLORS_HEX["border"],
                          width=1.5, alpha=0.7)
    nx.draw_networkx_nodes(G, pos, ax=ax, node_color=node_colors,
                          node_size=1500, alpha=0.9,
                          edgecolors="white", linewidths=2)
    nx.draw_networkx_labels(G, pos, labels=node_labels, ax=ax,
                           font_size=7, font_weight="bold")

    edge_labels = nx.get_edge_attributes(G, "label")
    nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, ax=ax,
                                font_size=6, font_color=COLORS_HEX["muted"])

    ax.set_title("Network Topology", fontsize=14, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=20)
    ax.axis("off")

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


def add_cloud_arch(doc, data: dict) -> None:
    """Embed the Diagram AI Generator cloud figure (PNG) when cloud infra is detected.

    No-op when there is no cloud infrastructure, or when graphviz/`diagrams` are
    unavailable — the rest of the report is unaffected.
    """
    if not isinstance(data, dict):
        return
    try:
        from cti_cloud_arch import build_cloud_png
    except Exception:
        return
    png, _ = build_cloud_png(data)
    if not png:
        return
    # pandoc-generated DOCX lacks the named "Heading 2" style, so add_heading(...)
    # by name raises KeyError. Set the style id in XML instead (matches the report's
    # Heading2 style and appears in the TOC), the same approach used elsewhere.
    heading = doc.add_paragraph("Cloud Architecture")
    heading._p.get_or_add_pPr().insert(
        0, parse_xml(f'<w:pStyle {nsdecls("w")} w:val="Heading2"/>'))
    doc.add_picture(BytesIO(png), width=Inches(5.6))
    doc.paragraphs[-1].alignment = 1
