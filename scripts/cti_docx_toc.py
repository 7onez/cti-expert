"""Self-contained Table of Contents for CTI DOCX reports.

The TOC is a real Word ``TOC`` field, so Word regenerates a proper page-numbered
table of contents when the document is opened (the generators also set
``<w:updateFields>`` in ``word/settings.xml`` to trigger that automatically).

Headless LibreOffice — the PDF path — never updates fields, and it cannot be made
to via ``--convert-to`` (and the minimal ``libreoffice-writer`` install ships no
pyuno bridge for a macro). So rather than leave the useless
``[Right-click and Update Field to generate TOC]`` placeholder as the field's
cached result, we bake a genuine, clickable heading list into the cache. Net
result:

* **Word** shows a full page-numbered TOC (field auto-updates on open).
* **LibreOffice / PDF / any un-updated view** shows a real clickable contents
  list — never the placeholder.

Two-pass usage (works for both the pandoc-hybrid and the pure python-docx
generators, since headings only exist after the body is built):

    anchor = begin_toc(doc)        # heading + invisible anchor at TOC position
    ...build the rest of the body...
    finalize_toc(doc, anchor)      # scan headings, bake clickable cached TOC
"""
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from cti_docx_styles import COLORS, FONT_HEADING, FONT_BODY, enable_update_fields

# Heading style ids emitted by both paths: python-docx built-ins ("Heading1") and
# pandoc ("Heading1"/"Heading 1"). Match by prefix, mirroring the postprocess helper.
_HEADING_PREFIXES = ("Heading", "heading")
_MAX_TOC_LEVEL = 3
_INDENT_TWIPS_PER_LEVEL = 360  # 0.25" per nesting level


def _heading_level(p_elem) -> int:
    """Heading level 1-9 for a ``w:p`` element, or 0 if it is not a heading."""
    if p_elem.tag != qn("w:p"):
        return 0
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return 0
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return 0
    val = pStyle.get(qn("w:val")) or ""
    for prefix in _HEADING_PREFIXES:
        if val.startswith(prefix):
            try:
                return int(val[len(prefix):].strip())
            except ValueError:
                pass
    return 0


def _text_of(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t"))).strip()


def _xml_escape(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
             .replace('"', "&quot;"))


def _hex(color) -> str:
    return f"{color:02X}" if isinstance(color, int) else str(color)


def begin_toc(doc):
    """Add the "Table of Contents" heading plus an invisible anchor paragraph at
    the current position. Returns the anchor paragraph; pass it to
    :func:`finalize_toc` after the rest of the body is built."""
    heading = doc.add_paragraph()
    run = heading.add_run("Table of Contents")
    run.font.bold = True
    run.font.size = _pt(18)
    run.font.color.rgb = COLORS["primary"]
    run.font.name = FONT_HEADING
    # Mark as Heading1 so it is styled consistently, but it sits BEFORE the anchor
    # so finalize_toc's scan (which starts after the anchor) never lists it.
    pPr = heading._element.get_or_add_pPr()
    pPr.insert(0, parse_xml(f'<w:pStyle {nsdecls("w")} w:val="Heading1"/>'))

    anchor = doc.add_paragraph()  # empty; the TOC field block is inserted after it
    return anchor


def _pt(points):
    from docx.shared import Pt
    return Pt(points)


def _collect_headings(doc, anchor):
    """Ordered (level, text, p_elem) for every Heading 1-3 that appears after the
    anchor in document order."""
    anchor_p = anchor._p
    body = anchor_p.getparent()
    started = False
    out = []
    for elem in body.iterchildren():
        if not started:
            if elem is anchor_p:
                started = True
            continue
        lvl = _heading_level(elem)
        if 1 <= lvl <= _MAX_TOC_LEVEL:
            text = _text_of(elem)
            if text:
                out.append((lvl, text, elem))
    return out


def _add_bookmark(p_elem, name: str, bid: int) -> None:
    start = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="{bid}" w:name="{_xml_escape(name)}"/>')
    end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bid}"/>')
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        p_elem.insert(0, start)
    p_elem.append(end)


def _opener_paragraph() -> "etree._Element":
    return parse_xml(
        f'<w:p {nsdecls("w")}>'
        '<w:pPr><w:spacing w:after="80"/></w:pPr>'
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        "</w:p>"
    )


def _entry_paragraph(level: int, text: str, bookmark: str) -> "etree._Element":
    left = (level - 1) * _INDENT_TWIPS_PER_LEVEL
    color = _hex(COLORS["steel"]) if level == 1 else _hex(COLORS["text"])
    sz = 22 if level == 1 else 20  # half-points -> 11pt / 10pt
    bold = "<w:b/>" if level == 1 else ""
    font = FONT_HEADING if level == 1 else FONT_BODY
    return parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr><w:ind w:left="{left}"/><w:spacing w:before="20" w:after="20"/></w:pPr>'
        f'<w:hyperlink w:anchor="{_xml_escape(bookmark)}">'
        f'<w:r><w:rPr>{bold}<w:rFonts w:ascii="{font}" w:hAnsi="{font}"/>'
        f'<w:color w:val="{color}"/><w:sz w:val="{sz}"/></w:rPr>'
        f'<w:t xml:space="preserve">{_xml_escape(text)}</w:t></w:r>'
        "</w:hyperlink>"
        "</w:p>"
    )


def _closer_paragraph() -> "etree._Element":
    return parse_xml(
        f'<w:p {nsdecls("w")}><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>')


def _page_break_paragraph() -> "etree._Element":
    return parse_xml(
        f'<w:p {nsdecls("w")}><w:r><w:br w:type="page"/></w:r></w:p>')


def finalize_toc(doc, anchor) -> int:
    """Scan headings after ``anchor``, bookmark them, and insert the TOC field
    block (opener + clickable cached entries + closer + page break) immediately
    after the anchor. Also flags document fields to refresh on open. Returns the
    number of TOC entries; if there are none, only the page break is added.
    """
    headings = _collect_headings(doc, anchor)

    block = [_opener_paragraph()]
    for i, (level, text, p_elem) in enumerate(headings, start=1):
        bm = f"_cti_toc_{i}"
        _add_bookmark(p_elem, bm, i)
        block.append(_entry_paragraph(level, text, bm))
    block.append(_closer_paragraph())
    block.append(_page_break_paragraph())

    cursor = anchor._p
    for elem in block:
        cursor.addnext(elem)
        cursor = elem

    enable_update_fields(doc)
    return len(headings)
