"""
CTI Report — Event Timeline figure (matplotlib -> PNG -> docx).

Split out of cti_docx_charts.py: the timeline carries its own layout engine (month grouping on a
date axis, label wrapping, per-row sizing, page-safe pagination) and was a third of that module.
cti_docx_charts re-exports add_timeline_chart, so callers are unchanged.
"""
import matplotlib
matplotlib.use("Agg")

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from io import BytesIO
from docx.shared import Inches

from cti_docx_styles import COLORS_HEX


def _save_fig_to_buffer(fig) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150, facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf


_TIMELINE_MAX_LABEL_LINES = 2
_TIMELINE_ROW_IN = 0.30       # inches per event row inside a month group
_TIMELINE_GROUP_GAP_IN = 0.22  # inches between month groups


def _parse_date(s):
    """YYYY-MM-DD / YYYY-MM / YYYY -> datetime.date, else None. Never raises."""
    import datetime as _dt
    s = (s or "").strip()[:10]
    for fmt in ("%Y-%m-%d", "%Y-%m", "%Y"):
        try:
            return _dt.datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def _common_suffix(strings: list) -> str:
    """Longest common word-aligned suffix shared by ALL strings (>=3 strings), else ''.
    Lets 'X registered through Global Domain Group LLC' x32 print the boilerplate once."""
    if len(strings) < 3:
        return ""
    words = [s.split() for s in strings]
    suffix = []
    for i in range(1, min(len(w) for w in words)):
        col = {w[-i] for w in words}
        if len(col) != 1:
            break
        suffix.insert(0, col.pop())
    return " ".join(suffix) if len(suffix) >= 2 else ""


def _wrap(text: str, width: int) -> str:
    import textwrap
    lines = textwrap.wrap(text, width=width, break_long_words=True, break_on_hyphens=False)
    if len(lines) > _TIMELINE_MAX_LABEL_LINES:
        lines = lines[:_TIMELINE_MAX_LABEL_LINES]
        lines[-1] = lines[-1][: max(0, width - 1)].rstrip() + "…"
    return "\n".join(lines) if lines else ""


def add_timeline_chart(doc, events: list) -> None:
    """Event timeline that stays legible from 3 to 100+ events.

    Dated events are grouped by month and plotted against a REAL date axis, so bursts
    (registration waves) read as clusters rather than as an undifferentiated list. Labels
    are wrapped, never truncated; boilerplate shared by every event is factored into one
    caption. Undated / unparseable events fall back to a compact two-column list.
    """
    if not events:
        return

    rows = []
    for e in events:
        label = (e.get("event") or "").strip()
        d = _parse_date(e.get("date"))
        rows.append((d, e.get("date") or "N/A", label))
    dated = [r for r in rows if r[0] is not None]
    undated = [r for r in rows if r[0] is None]

    labels_all = [r[2] for r in rows if r[2]]
    suffix = _common_suffix(labels_all)
    def _short(label: str) -> str:
        if suffix and label.endswith(suffix):
            label = label[: -len(suffix)].rstrip(" ,;:-—–")
        return label or "(event)"

    if len(dated) >= 3:
        _timeline_by_month(doc, sorted(dated, key=lambda r: r[0]), undated, suffix, _short)
    else:
        _timeline_list(doc, sorted(dated, key=lambda r: r[0]) + undated, suffix, _short)


def _timeline_by_month(doc, dated, undated, suffix, _short) -> None:
    """Month-grouped swimlane: one band per month with its events listed, on a date axis."""
    from collections import OrderedDict
    groups = OrderedDict()
    for d, raw, label in dated:
        groups.setdefault((d.year, d.month), []).append((d, _short(label)))

    n_events = len(dated)
    n_groups = len(groups)
    two_col = n_events > 16
    label_width = 31 if two_col else 80
    fig_w = 7.2
    ncols = 2 if two_col else 1
    _LINE_IN = 0.175   # vertical inches per wrapped text line at 7.5 pt
    _ROW_PAD_IN = 0.16
    _HEADER_IN = 0.36

    # Pre-wrap every label and size each row by the tallest label in it, so a two-line
    # label never collides with the row below it.
    laid = []  # per group: (rows=[(row_h, [(d, wrapped), ...]), ...], band_h)
    for items in groups.values():
        wrapped = [(d, _wrap(lbl, label_width)) for d, lbl in items]
        rows_ = []
        for i in range(0, len(wrapped), ncols):
            chunk = wrapped[i:i + ncols]
            n_lines = max(w.count("\n") + 1 for _, w in chunk)
            rows_.append((_ROW_PAD_IN + _LINE_IN * n_lines, chunk))
        laid.append((rows_, _HEADER_IN + sum(h for h, _ in rows_)))

    # Paginate: a DOCX picture cannot break across pages, so keep each figure under one page
    # of usable height (6.3 in wide placement scales 7.2 in → ×0.875; 9.6 in data ≈ 8.4 in on page).
    _PAGE_IN = 9.6
    # 1) split any single month that is itself taller than a page (a one-month registration burst)
    units = []   # (key, items, (rows_, band_h), cont_flag)
    for gi, key in enumerate(groups.keys()):
        rows_, band_h = laid[gi]
        if band_h + _TIMELINE_GROUP_GAP_IN <= _PAGE_IN:
            units.append((key, groups[key], (rows_, band_h), False))
            continue
        chunk, ch_h, first = [], 0.0, True
        for row in rows_:
            if chunk and ch_h + row[0] + _HEADER_IN + _TIMELINE_GROUP_GAP_IN > _PAGE_IN:
                units.append((key, [d for _, c_ in chunk for d, _ in c_], (chunk, _HEADER_IN + ch_h), not first))
                chunk, ch_h, first = [], 0.0, False
            chunk.append(row); ch_h += row[0]
        if chunk:
            units.append((key, [d for _, c_ in chunk for d, _ in c_], (chunk, _HEADER_IN + ch_h), not first))
    # 2) first-fit to learn how many pages are needed …
    heights = [u[2][1] + _TIMELINE_GROUP_GAP_IN for u in units]
    pages, cur, cur_h = [], [], 0.0
    for i, h in enumerate(heights):
        if cur and cur_h + h > _PAGE_IN:
            pages.append(cur); cur, cur_h = [], 0.0
        cur.append(i); cur_h += h
    if cur:
        pages.append(cur)
    # 3) … then re-partition greedily against an even target so pages are balanced (5+4, not 8+1)
    if len(pages) > 1:
        n_pages = len(pages)
        target = sum(heights) / n_pages
        pages, cur, cur_h = [], [], 0.0
        for i, h in enumerate(heights):
            if cur and (cur_h + h > _PAGE_IN or (cur_h + h / 2 > target and len(pages) < n_pages - 1)):
                pages.append(cur); cur, cur_h = [], 0.0
            cur.append(i); cur_h += h
        if cur:
            pages.append(cur)

    first, last = dated[0][0], dated[-1][0]
    span = f"{first.isoformat()} → {last.isoformat()}  ·  {n_events} events in {n_groups} month{'s' if n_groups != 1 else ''}"
    caption = []
    if suffix:
        caption.append(f"All events: “… {suffix}”")
    if undated:
        caption.append(f"{len(undated)} undated event{'s' if len(undated) != 1 else ''} listed below")

    for pi, page in enumerate(pages):
        title = "Event Timeline" + (f"  ({pi + 1}/{len(pages)})" if len(pages) > 1 else "")
        _timeline_page(doc, [units[i] for i in page], fig_w, title,
                       span if pi == 0 else f"continued · {span}", "  ·  ".join(caption),
                       cont_before=pi > 0, cont_after=pi < len(pages) - 1)

    if undated:
        _timeline_list(doc, undated, "", _short, title="Undated events")


def _timeline_page(doc, page_groups, fig_w, title, span, caption, cont_before, cont_after) -> None:
    """Render one page of month bands. page_groups: [((yr, mo), items, (rows_, band_h), cont), ...]."""
    _HEADER_IN = 0.36
    total_h = sum(band_h for _, _, (_, band_h), _c in page_groups) + len(page_groups) * _TIMELINE_GROUP_GAP_IN
    fig_h = max(2.4, total_h + 0.75)  # 1 data unit == 1 inch: nothing is squeezed

    fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=150)
    fig.subplots_adjust(left=0.02, right=0.98, top=1 - 0.45 / fig_h, bottom=0.25 / fig_h)
    ax.set_xlim(0, 10)
    ax.axis("off")
    ax.set_ylim(total_h, -0.05)  # inverted: earliest at top

    month_names = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    spine_x = 0.55
    y = 0.0
    node_ys = []
    for (yr, mo), items, (rows_, band_h), cont in page_groups:
        # Month band
        ax.add_patch(mpatches.FancyBboxPatch(
            (1.05, y), 8.9, band_h, boxstyle="round,pad=0,rounding_size=0.08",
            linewidth=0, facecolor=COLORS_HEX["bg_light"], zorder=1))
        # Node on the spine, sized by burst size
        node_ys.append(y + 0.18)
        ax.scatter([spine_x], [y + 0.18], s=90 + 18 * len(items), c=COLORS_HEX["accent"],
                   zorder=5, edgecolors="white", linewidths=1.5)
        ax.text(spine_x, y + 0.18, str(len(items)), ha="center", va="center", fontsize=7,
                color="white", fontweight="bold", zorder=6)
        # Month header
        ax.text(1.2, y + 0.16, f"{month_names[mo - 1]} {yr}" + (" (cont.)" if cont else ""), va="center", ha="left",
                fontsize=10.5, fontweight="bold", color=COLORS_HEX["primary"], zorder=3)
        ax.text(9.85, y + 0.16, f"{len(items)} event{'s' if len(items) != 1 else ''}",
                va="center", ha="right", fontsize=8, color=COLORS_HEX["muted"], zorder=3)
        # Events
        yy = y + _HEADER_IN
        for row_h, chunk in rows_:
            for col, (d, wrapped) in enumerate(chunk):
                x0 = 1.25 + col * 4.45
                ax.text(x0, yy + 0.06, d.strftime("%d"), va="top", ha="left", fontsize=7.5,
                        color=COLORS_HEX["ochre"], fontweight="bold", family="monospace", zorder=3)
                ax.text(x0 + 0.42, yy + 0.06, wrapped, va="top", ha="left",
                        fontsize=8.5, color=COLORS_HEX["text"], zorder=3, linespacing=1.15)
            yy += row_h
        y += band_h + _TIMELINE_GROUP_GAP_IN

    # Spine through the month nodes; dashed stubs show continuation across pages.
    if node_ys:
        y0 = -0.05 if cont_before else node_ys[0]
        y1 = total_h if cont_after else node_ys[-1]
        ax.vlines(spine_x, y0, y1, color=COLORS_HEX["border"], linewidth=2, zorder=1,
                  linestyles="dashed" if (cont_before or cont_after) else "solid")

    ax.set_title(title, fontsize=12, fontweight="bold", color=COLORS_HEX["primary"], pad=18, loc="left")
    ax.text(0, 1.0, span, transform=ax.transAxes, ha="left", va="bottom", fontsize=8, color=COLORS_HEX["muted"])
    if caption:
        fig.text(0.01, 0.005, caption, ha="left", va="bottom", fontsize=7.5,
                 color=COLORS_HEX["muted"], style="italic")

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(6.3))
    doc.paragraphs[-1].alignment = 1


_LIST_ROWS_PER_PAGE = 36   # two columns x 18 rows ≈ 8 in of figure


def _timeline_list(doc, rows, suffix, _short, title: str = "Event Timeline") -> None:
    """Compact fallback for few or undated events: two-column dot list, full text, wrapped.
    Long lists paginate (one picture per _LIST_ROWS_PER_PAGE rows) instead of clamping."""
    if not rows:
        return
    if len(rows) > _LIST_ROWS_PER_PAGE:
        n = (len(rows) + _LIST_ROWS_PER_PAGE - 1) // _LIST_ROWS_PER_PAGE
        for i in range(n):
            _timeline_list(doc, rows[i * _LIST_ROWS_PER_PAGE:(i + 1) * _LIST_ROWS_PER_PAGE], suffix if i == 0 else "",
                           _short, title=f"{title}  ({i + 1}/{n})")
        return
    two_col = len(rows) > 8
    per_col = (len(rows) + 1) // 2 if two_col else len(rows)
    label_width = 33 if two_col else 58
    fig_h = max(1.6, 0.7 + per_col * 0.46)
    fig, ax = plt.subplots(figsize=(7.2, fig_h), dpi=150)
    fig.subplots_adjust(left=0.02, right=0.98, top=1 - 0.4 / fig_h, bottom=0.25 / fig_h)
    ax.set_xlim(0, 10)
    ax.set_ylim(per_col, -0.6)
    ax.axis("off")
    for i, (d, raw, label) in enumerate(rows):
        col = i // per_col if two_col else 0
        r = i % per_col if two_col else i
        x0 = 0.35 + col * 5.0
        ax.scatter([x0], [r], s=60, c=COLORS_HEX["accent"], zorder=5, edgecolors="white", linewidths=1.2)
        ax.text(x0 + 0.3, r, raw[:10], va="center", ha="left", fontsize=7.5,
                color=COLORS_HEX["ochre"], fontweight="bold", family="monospace")
        ax.text(x0 + 1.55, r, _wrap(_short(label), label_width), va="center", ha="left",
                fontsize=8.5, color=COLORS_HEX["text"], linespacing=1.15)
        if r < per_col - 1 and (i + 1 < len(rows)) and ((i + 1) // per_col == col if two_col else True):
            ax.vlines(x0, r, r + 1, color=COLORS_HEX["border"], linewidth=1.5, zorder=1)
    ax.set_title(title, fontsize=12, fontweight="bold", color=COLORS_HEX["primary"], pad=12, loc="left")
    if suffix:
        fig.text(0.01, 0.005, f"All events: “… {suffix}”", ha="left", va="bottom", fontsize=7.5,
                 color=COLORS_HEX["muted"], style="italic")
    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(6.3))
    doc.paragraphs[-1].alignment = 1


