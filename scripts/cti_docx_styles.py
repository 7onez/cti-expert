"""
CTI Report DOCX Styles — fonts, colors, heading styles, header/footer.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# --- CTI House Palette ---------------------------------------------------------
# Muted editorial palette, IDENTICAL to the PDF (IntelGraph/scripts/theme.py +
# IntelReport/templates/house-header.tex): slate / steel / ochre / brick / grid.
# So a report's DOCX reads as the same document as its PDF, not a different tool's
# output. Change a colour HERE and every DOCX chart, table and heading follows.
FONT_BODY = "Georgia"      # serif body — the house choice (read in print, at length)
FONT_HEADING = "Arial"     # sans headings / furniture — distinct from the serif body
FONT_MONO = "Consolas"     # inline code: domains, hashes, endpoints

COLORS = {
    "primary": RGBColor(0x22, 0x33, 0x3F),       # slate — H1, cover title, strong
    "steel":   RGBColor(0x3B, 0x55, 0x66),       # steel — H2/H3, links
    "accent":  RGBColor(0x3B, 0x55, 0x66),       # steel (alias kept for callers)
    "ochre":   RGBColor(0xB0, 0x79, 0x0F),
    "brick":   RGBColor(0x8C, 0x2D, 0x2D),
    "olive":   RGBColor(0x5A, 0x6B, 0x3B),
    "critical": RGBColor(0x8C, 0x2D, 0x2D),       # brick
    "high":    RGBColor(0xB0, 0x79, 0x0F),        # ochre
    "medium":  RGBColor(0x9A, 0x5B, 0x2F),        # muted amber-brown
    "low":     RGBColor(0x5A, 0x6B, 0x3B),        # olive
    "info":    RGBColor(0x3B, 0x55, 0x66),        # steel
    "text":    RGBColor(0x1F, 0x1D, 0x1A),        # ink
    "muted":   RGBColor(0x6F, 0x6A, 0x61),        # muted
    "white":   RGBColor(0xFF, 0xFF, 0xFF),
    "band":    RGBColor(0xDE, 0xE4, 0xE8),        # table header band (light steel-grey)
    "bg_light": RGBColor(0xF2, 0xF2, 0xF2),       # zebra stripe / subtle fill
    "callout": RGBColor(0xF7, 0xF4, 0xED),        # IC callout background
    "border":  RGBColor(0xD9, 0xD3, 0xC7),        # grid
}

# Hex versions for matplotlib (charts) — same palette.
COLORS_HEX = {
    "primary": "#22333F",   # slate
    "steel":   "#3B5566",
    "accent":  "#3B5566",   # steel
    "ochre":   "#B0790F",
    "brick":   "#8C2D2D",
    "olive":   "#5A6B3B",
    "critical": "#8C2D2D",
    "high":    "#B0790F",
    "medium":  "#9A5B2F",
    "low":     "#5A6B3B",
    "info":    "#3B5566",
    "text":    "#1F1D1A",
    "muted":   "#6F6A61",
    "band":    "#DEE4E8",
    "bg_light": "#F2F2F2",
    "border":  "#D9D3C7",
}

# Ordered categorical cycle for charts (colorblind-safe, no default matplotlib
# blue) — matches theme.py CYCLE so a DOCX pie and a PDF figure agree.
CYCLE_HEX = ["#3B5566", "#8C2D2D", "#B0790F", "#5A6B3B", "#22333F", "#C9B892",
             "#5A4A7A", "#2F6B6B"]

SEVERITY_COLORS_HEX = {
    "CRITICAL": "#8C2D2D",
    "HIGH": "#B0790F",
    "MEDIUM": "#9A5B2F",
    "LOW": "#5A6B3B",
    "INFO": "#3B5566",
}

SEVERITY_COLORS = {
    "CRITICAL": COLORS["critical"],
    "HIGH": COLORS["high"],
    "MEDIUM": COLORS["medium"],
    "LOW": COLORS["low"],
    "INFO": COLORS["info"],
}

# Cell-shading strings (raw hex, no '#') DERIVED from the palette above — so the
# one place to change the header band / zebra colour is COLORS_HEX, not scattered
# literals in postprocess/sections. set_cell_shading() wants the bare hex.
SHADE_BAND = COLORS_HEX["band"].lstrip("#")       # table header band
SHADE_ZEBRA = COLORS_HEX["bg_light"].lstrip("#")  # zebra stripe


def setup_styles(doc: Document) -> Document:
    """Configure document styles for CTI report."""
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLORS["text"]
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size, color in [
        ("Heading 1", 18, "primary"),   # slate
        ("Heading 2", 14, "steel"),
        ("Heading 3", 12, "steel"),
    ]:
        s = doc.styles[level]
        s.font.name = FONT_HEADING
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = COLORS[color]
        s.paragraph_format.space_before = Pt(18 if level == "Heading 1" else 12)
        s.paragraph_format.space_after = Pt(8)

    return doc


def add_page_number(paragraph):
    """Add page number field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def setup_header_footer(doc: Document, report_id: str, classification: str = "OPEN SOURCE"):
    """Add header with classification and footer with page numbers."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(f"CTI REPORT  |  {classification}  |  {report_id}")
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["muted"]
    run.font.name = FONT_HEADING

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"CTI Report — {report_id}  |  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["muted"]
    run.font.name = FONT_HEADING
    add_page_number(fp)

    # Credit line — CTI Expert attribution on every generated report.
    cp = footer.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cp.add_run("Generated by CTI Expert  ·  https://github.com/7onez/cti-expert")
    crun.font.size = Pt(7)
    crun.font.color.rgb = COLORS["muted"]
    crun.font.name = FONT_HEADING

    return doc


def add_cover_page(doc: Document, data: dict) -> Document:
    """Create professional cover page with CTI Report title."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CTI REPORT")
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = FONT_HEADING

    # Subtitle — case label
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(data.get("case", {}).get("label", "Intelligence Summary"))
    run.font.size = Pt(18)
    run.font.color.rgb = COLORS["accent"]
    run.font.name = FONT_HEADING

    # Divider line
    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div_para.add_run("_" * 50)
    run.font.color.rgb = COLORS["border"]

    doc.add_paragraph()

    # Metadata table
    case = data.get("case", {})
    meta_items = [
        ("Report ID", case.get("id", "N/A")),
        ("Classification", case.get("classification", "OPEN SOURCE")),
        ("Date", case.get("date", datetime.date.today().isoformat())),
        ("Analyst", case.get("analyst", "AI-Assisted OSINT")),
        ("Subject", case.get("subject", "N/A")),
        ("Status", case.get("status", "active")),
    ]

    table = doc.add_table(rows=len(meta_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta_items):
        row = table.rows[i]
        cell_l = row.cells[0]
        cell_r = row.cells[1]
        cell_l.width = Inches(2)
        cell_r.width = Inches(4)

        run_l = cell_l.paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.size = Pt(11)
        run_l.font.color.rgb = COLORS["muted"]
        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run_r = cell_r.paragraphs[0].add_run(value)
        run_r.font.size = Pt(11)
        run_r.font.color.rgb = COLORS["text"]

    # Remove table borders for clean look
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                "</w:tcBorders>"
            )
            tcPr.append(tcBorders)

    # Page break after cover
    doc.add_page_break()

    return doc


def add_table_of_contents(doc: Document) -> Document:
    """Add a TOC field — requires Word to refresh on open."""
    toc_heading = doc.add_paragraph("Table of Contents", style="Heading 1")

    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run("[Right-click and Update Field to generate TOC]")
    run4.font.color.rgb = COLORS["muted"]
    run4.font.size = Pt(9)
    run4.font.italic = True

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    doc.add_page_break()
    return doc


def set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)
