"""
Post-processing for pandoc-generated DOCX files.
Applies CTI professional styling, injects charts from JSON data,
and prepends cover page + table of contents.
"""
import sys
import unicodedata

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import datetime
import base64
from io import BytesIO

from cti_docx_styles import (
    COLORS, set_cell_shading, add_page_number,
    FONT_BODY, FONT_HEADING, SHADE_BAND, SHADE_ZEBRA,
)
from cti_docx_charts import (
    add_finding_type_pie, add_severity_bar, add_confidence_matrix,
    add_timeline_chart, add_traffic_sources_bar, add_geographic_pie,
)
from cti_docx_diagrams import add_entity_diagram, add_network_topology, add_cloud_arch
from cti_docx_heatmaps import (
    collect_registrations, add_registration_heatmap,
    build_cooccurrence, add_cooccurrence_heatmap,
    build_relation_strength, add_relation_heatmap,
)


def add_evidence_screenshots(doc, images) -> None:
    """Embed evidence screenshots (base64 data URIs) as centered figures with captions.
    Each image dict: {caption?, host?, data_uri, source_url?, captured_at?, finding_id?}.
    A bad/undecodable entry is skipped so one broken image never breaks the report."""
    if not images:
        return
    for im in images:
        if not isinstance(im, dict):
            continue
        uri = im.get("data_uri") or ""
        b64 = uri.split(",", 1)[1] if (uri.startswith("data:") and "," in uri) else uri
        if not b64:
            continue
        try:
            raw = base64.b64decode(b64)
            doc.add_picture(BytesIO(raw), width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:
            sys.stderr.write("evidence screenshot skipped (%s): %s\n"
                             % (im.get("caption") or im.get("host") or "?", exc))
            continue
        cap_bits = [b for b in (
            im.get("caption"), im.get("host"),
            ("finding " + str(im.get("finding_id"))) if im.get("finding_id") else None,
            im.get("captured_at"),
        ) if b]
        if cap_bits:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" · ".join(str(b) for b in cap_bits))
            run.font.size = Pt(9)
            run.font.italic = True
        doc.add_paragraph()


CHART_KEYWORDS = {
    "confidence_matrix": ["confidence", "assessment", "danh gia", "do tin cay", "icd-203", "admiralty", "nhan dinh", "key judgment"],
    "finding_charts": ["phat hien", "findings", "statistical"],
    "timeline_chart": ["timeline", "thoi gian", "dong thoi gian"],
    "entity_diagram": ["moi quan he", "relationship", "entity", "ban do"],
    "heatmaps": ["correlation", "tuong quan", "registration", "dang ky", "lien ket", "cluster", "campaign", "chien dich", "infrastructure", "ha tang"],
    "visitor_charts": ["visitor", "traffic", "luong truy cap", "hien dien"],
    "evidence_screenshots": ["evidence screenshot", "screenshot", "visual evidence", "anh chup man hinh", "bang chung hinh anh"],
}


def strip_accents(text: str) -> str:
    nfkd = unicodedata.normalize("NFKD", text)
    return "".join(c for c in nfkd if not unicodedata.combining(c))


def _heading_matches(text: str, keywords: list[str]) -> bool:
    normalized = strip_accents(text.lower().strip())
    return any(kw in normalized for kw in keywords)


def _extract_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(qn("w:t")))


def _get_heading_level_from_xml(elem) -> int:
    """Return heading level (1-9) from a w:p element, or 0 if not a heading."""
    if elem.tag != qn("w:p"):
        return 0
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return 0
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return 0
    val = pStyle.get(qn("w:val")) or ""
    for prefix in ("Heading", "heading"):
        if val.startswith(prefix):
            try:
                return int(val[len(prefix):].strip())
            except ValueError:
                pass
    return 0


def _add_styled_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading paragraph using XML style IDs, bypassing python-docx name lookup
    which fails on pandoc-generated documents."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    size, color_key = HEADING_THEME.get(level, (12, "accent"))
    run.font.size = Pt(size)
    run.font.color.rgb = COLORS[color_key]
    run.font.name = FONT_HEADING
    pPr = p._element.get_or_add_pPr()
    existing_pStyle = pPr.find(qn("w:pStyle"))
    if existing_pStyle is not None:
        pPr.remove(existing_pStyle)
    pPr.insert(0, parse_xml(f'<w:pStyle {nsdecls("w")} w:val="Heading{level}"/>'))


def _add_cover_page_compat(doc: Document, data: dict) -> None:
    """Cover page that works with pandoc-generated DOCX (no style name lookups)."""
    from docx.shared import Inches, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT

    for _ in range(4):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CTI REPORT")
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = FONT_HEADING

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(data.get("case", {}).get("label", "Intelligence Summary"))
    run.font.size = Pt(18)
    run.font.color.rgb = COLORS["accent"]
    run.font.name = FONT_HEADING

    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div_para.add_run("_" * 50)
    run.font.color.rgb = COLORS["border"]

    doc.add_paragraph()

    case = data.get("case", {})
    meta_items = [
        ("Report ID", case.get("id", "N/A")),
        ("Classification", case.get("classification", "OPEN SOURCE")),
        ("Date", case.get("date", datetime.date.today().isoformat())),
        ("Analyst", case.get("analyst", "AI-Assisted OSINT")),
        ("Subject", case.get("subject", "N/A")),
        ("Status", case.get("status", "active")),
    ]

    table = doc.add_table(rows=len(meta_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta_items):
        row = table.rows[i]
        cell_l, cell_r = row.cells[0], row.cells[1]
        cell_l.width = Inches(2)
        cell_r.width = Inches(4)
        run_l = cell_l.paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.size = Pt(11)
        run_l.font.color.rgb = COLORS["muted"]
        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_r = cell_r.paragraphs[0].add_run(value)
        run_r.font.size = Pt(11)
        run_r.font.color.rgb = COLORS["text"]

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                "</w:tcBorders>"
            ))

    doc.add_page_break()


def _add_toc_compat(doc: Document) -> None:
    """Table of contents that works with pandoc-generated DOCX."""
    _add_styled_heading(doc, "Table of Contents", level=1)

    p = doc.add_paragraph()
    run = p.add_run()
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run2 = p.add_run()
    run2._r.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    ))
    run3 = p.add_run()
    run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    run4 = p.add_run("[Right-click and Update Field to generate TOC]")
    run4.font.color.rgb = COLORS["muted"]
    run4.font.size = Pt(9)
    run4.font.italic = True
    run5 = p.add_run()
    run5._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    doc.add_page_break()


def setup_header_footer_compat(doc: Document, report_id: str,
                                classification: str = "OPEN SOURCE") -> None:
    """Header/footer that works with pandoc-generated DOCX."""
    from docx.shared import Cm
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(f"CTI REPORT  |  {classification}  |  {report_id}")
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["muted"]
    run.font.name = FONT_HEADING

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"CTI Report — {report_id}  |  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["muted"]
    run.font.name = FONT_HEADING
    add_page_number(fp)

    # Credit line — CTI Expert attribution on every generated report.
    cp = footer.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cp.add_run("Generated by CTI Expert  ·  https://github.com/7onez/cti-expert")
    crun.font.size = Pt(7)
    crun.font.color.rgb = COLORS["muted"]
    crun.font.name = FONT_HEADING


def apply_cti_styles(doc: Document) -> None:
    """Restyle a pandoc-generated DOCX with CTI theme colors and fonts.

    pandoc-generated DOCX styles can't be accessed via doc.styles["Heading 1"]
    (KeyError), but are accessible via iteration. We iterate all styles and
    apply CTI theming to any heading or body style found.
    """
    _apply_heading_styles_by_iteration(doc)
    _set_normal_style(doc)
    for table in doc.tables:
        _style_table(table)


HEADING_THEME = {
    1: (18, "primary"),   # slate
    2: (14, "steel"),
    3: (12, "steel"),
}


def _apply_heading_styles_by_iteration(doc: Document) -> None:
    for style in doc.styles:
        name = style.name or ""
        if not name.startswith("Heading "):
            continue
        try:
            level = int(name.split()[-1])
        except (ValueError, IndexError):
            continue
        size, color_key = HEADING_THEME.get(level, (12, "steel"))
        style.font.name = FONT_HEADING
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLORS[color_key]
        if style.paragraph_format:
            style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            style.paragraph_format.space_after = Pt(8)


def _set_normal_style(doc: Document) -> None:
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        return
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = COLORS["text"]
    if normal.paragraph_format:
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15


def _style_table(table) -> None:
    if not table.rows:
        return
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for cell in table.rows[0].cells:
        set_cell_shading(cell, SHADE_BAND)   # light steel header band
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = COLORS["text"]   # dark ink on light band
                run.font.bold = True
                run.font.size = Pt(9)
    # Zebra: shade alternate BODY rows (subtle neutral grey) for legibility.
    for i, row in enumerate(table.rows[1:]):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_shading(cell, SHADE_ZEBRA)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)


def rebuild_with_cover_toc_and_charts(doc: Document, json_data: dict, case_dir: str = None) -> None:
    """Remove all body content, prepend cover + TOC, re-add content with
    charts injected at matching section headings, then append remaining charts."""
    body = doc.element.body
    existing = list(body)

    sect_pr = None
    content_elements = []
    for elem in existing:
        if elem.tag == qn("w:sectPr"):
            sect_pr = elem
        else:
            content_elements.append(elem)

    for elem in existing:
        body.remove(elem)


    # Keep the section properties as the final body child so python-docx section
    # access works during cover/TOC; every subsequent add (cover, TOC, narrative,
    # inline charts, appendix) is inserted BEFORE it, preserving document order.
    if sect_pr is not None:
        body.append(sect_pr)

    _add_cover_page_compat(doc, json_data)
    _add_toc_compat(doc)

    injected: set[str] = set()
    case = json_data.get("case", {})

    for elem in content_elements:
        if sect_pr is not None:
            sect_pr.addprevious(elem)
        else:
            body.append(elem)

        level = _get_heading_level_from_xml(elem)
        if level == 0:
            continue
        text = _extract_text(elem)
        _try_inject(doc, text, json_data, injected, case, case_dir)

    _append_remaining_charts(doc, json_data, injected, case, case_dir)

    if sect_pr is not None:
        body.append(sect_pr)


def _try_inject(doc, heading_text, json_data, injected, case, case_dir=None):
    if "confidence_matrix" not in injected and _heading_matches(heading_text, CHART_KEYWORDS["confidence_matrix"]):
        findings = json_data.get("findings", [])
        if findings:
            doc.add_paragraph()
            add_confidence_matrix(doc, findings, json_data.get("subjects", []))
            doc.add_paragraph()
            injected.add("confidence_matrix")

    if "heatmaps" not in injected and _heading_matches(heading_text, CHART_KEYWORDS["heatmaps"]):
        if _inject_heatmaps(doc, json_data, case_dir):
            injected.add("heatmaps")

    if "finding_charts" not in injected and _heading_matches(heading_text, CHART_KEYWORDS["finding_charts"]):
        findings = json_data.get("findings", [])
        if findings:
            doc.add_paragraph()
            add_finding_type_pie(doc, findings)
            doc.add_paragraph()
            add_severity_bar(doc, findings)
            doc.add_paragraph()
            injected.add("finding_charts")

    if "timeline_chart" not in injected and _heading_matches(heading_text, CHART_KEYWORDS["timeline_chart"]):
        timeline = json_data.get("timeline", [])
        if timeline:
            doc.add_paragraph()
            add_timeline_chart(doc, timeline)
            doc.add_paragraph()
            injected.add("timeline_chart")

    if "entity_diagram" not in injected and _heading_matches(heading_text, CHART_KEYWORDS["entity_diagram"]):
        subjects = json_data.get("subjects", [])
        connections = json_data.get("connections", [])
        if subjects and connections:
            doc.add_paragraph()
            add_entity_diagram(doc, subjects, connections, json_data)
            doc.add_paragraph()
            add_network_topology(doc, subjects, connections, json_data)
            add_cloud_arch(doc, json_data)
            doc.add_paragraph()
            injected.add("entity_diagram")

    if "visitor_charts" not in injected and _heading_matches(heading_text, CHART_KEYWORDS["visitor_charts"]):
        vs = json_data.get("visitor_stats", {})
        if vs:
            ts = vs.get("traffic_sources", {})
            if ts:
                doc.add_paragraph()
                add_traffic_sources_bar(doc, ts)
            tc = vs.get("top_countries", [])
            if tc:
                doc.add_paragraph()
                add_geographic_pie(doc, tc)
            doc.add_paragraph()
            injected.add("visitor_charts")
    if "evidence_screenshots" not in injected and _heading_matches(heading_text, CHART_KEYWORDS["evidence_screenshots"]):
        images = json_data.get("evidence_images", [])
        if images:
            doc.add_paragraph()
            add_evidence_screenshots(doc, images)
            injected.add("evidence_screenshots")


def _inject_heatmaps(doc, json_data, case_dir=None):
    """Render the three campaign heatmaps; return True if at least one was drawn."""
    drawn = False
    if add_registration_heatmap(doc, collect_registrations(json_data, case_dir)):
        doc.add_paragraph(); drawn = True
    if add_cooccurrence_heatmap(doc, build_cooccurrence(json_data, case_dir)):
        doc.add_paragraph(); drawn = True
    if add_relation_heatmap(doc, build_relation_strength(json_data, case_dir)):
        doc.add_paragraph(); drawn = True
    return drawn


def _append_remaining_charts(doc, json_data, injected, case, case_dir=None):
    """Append any charts that were not injected inline as a Visual Analytics appendix."""
    remaining = set(CHART_KEYWORDS.keys()) - injected
    if not remaining:
        return

    has_content = False

    if "confidence_matrix" in remaining:
        findings = json_data.get("findings", [])
        if findings:
            if not has_content:
                doc.add_page_break()
                _add_styled_heading(doc, "Visual Analytics", level=1)
                has_content = True
            _add_styled_heading(doc, "Confidence Scale (ICD-203 \u00d7 Admiralty)", level=2)
            add_confidence_matrix(doc, findings, json_data.get("subjects", []))
            doc.add_paragraph()

    if "heatmaps" in remaining:
        if not has_content:
            doc.add_page_break()
            _add_styled_heading(doc, "Visual Analytics", level=1)
            has_content = True
        _add_styled_heading(doc, "Campaign Heatmaps", level=2)
        _inject_heatmaps(doc, json_data, case_dir)

    if "finding_charts" in remaining:
        findings = json_data.get("findings", [])
        if findings:
            if not has_content:
                doc.add_page_break()
                _add_styled_heading(doc, "Visual Analytics", level=1)
                has_content = True
            _add_styled_heading(doc, "Finding Statistics", level=2)
            add_finding_type_pie(doc, findings)
            doc.add_paragraph()
            add_severity_bar(doc, findings)
            doc.add_paragraph()

    if "timeline_chart" in remaining:
        timeline = json_data.get("timeline", [])
        if timeline:
            if not has_content:
                doc.add_page_break()
                _add_styled_heading(doc, "Visual Analytics", level=1)
                has_content = True
            _add_styled_heading(doc, "Event Timeline", level=2)
            add_timeline_chart(doc, timeline)
            doc.add_paragraph()

    if "entity_diagram" in remaining:
        subjects = json_data.get("subjects", [])
        connections = json_data.get("connections", [])
        if subjects and connections:
            if not has_content:
                doc.add_page_break()
                _add_styled_heading(doc, "Visual Analytics", level=1)
                has_content = True
            _add_styled_heading(doc, "Entity Relationships", level=2)
            add_entity_diagram(doc, subjects, connections, json_data)
            doc.add_paragraph()
            _add_styled_heading(doc, "Network Topology", level=2)
            add_network_topology(doc, subjects, connections, json_data)
            add_cloud_arch(doc, json_data)
            doc.add_paragraph()

    if "visitor_charts" in remaining:
        vs = json_data.get("visitor_stats", {})
        if vs:
            if not has_content:
                doc.add_page_break()
                _add_styled_heading(doc, "Visual Analytics", level=1)
                has_content = True
            _add_styled_heading(doc, "Visitor Intelligence", level=2)
            ts = vs.get("traffic_sources", {})
            if ts:
                add_traffic_sources_bar(doc, ts)
                doc.add_paragraph()
            tc = vs.get("top_countries", [])
            if tc:
                add_geographic_pie(doc, tc)
                doc.add_paragraph()

    if "evidence_screenshots" in remaining:
        images = json_data.get("evidence_images", [])
        if images:
            if not has_content:
                doc.add_page_break()
                _add_styled_heading(doc, "Visual Analytics", level=1)
                has_content = True
            _add_styled_heading(doc, "Evidence Screenshots", level=2)
            add_evidence_screenshots(doc, images)
