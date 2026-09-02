"""
CTI Report Heatmaps — malicious-domain registration timeline, domain x indicator
correlation (co-occurrence), and domain x domain possible-relation strength.

All builders read the report JSON primarily (subjects / connections / timeline) and
opportunistically enrich from case sidecars (whois/*.json created dates, case_graph.json
indicator edges) when a case_dir is supplied. Every builder degrades gracefully: no data
-> returns None and the caller skips the figure (never raises).
"""
import matplotlib
matplotlib.use("Agg")

import os
import re
import glob
import json
import datetime as _dt
from io import BytesIO

import numpy as np
import matplotlib.pyplot as plt

from cti_palette import COLORS_HEX

# Roles that mark a domain as NOT the operator's own infrastructure — excluded from
# the "malicious domains" registration timeline (the impersonated brand, confirmed
# victims, indicators already ruled benign).
_NON_MALICIOUS_ROLES = {"victim", "legitimate", "legit", "benign", "witness"}

# to_id namespaces in report connections that identify a SHARED registration attribute.
_ATTR_RELATIONSHIPS = {"registrant", "registrant_email", "registrar", "nameserver", "ns"}


DOMAIN_CAP = 24   # dashboard readability cap; the house report (Rule 19: no sampling) passes cap=None


def _domains(json_data, cap=DOMAIN_CAP):
    """Ordered list of domain subject labels (dedup, capped for readability unless cap=None)."""
    seen, out = set(), []
    for s in json_data.get("subjects", []) or []:
        if str(s.get("type", "")).lower() != "domain":
            continue
        label = (s.get("label") or "").strip()
        if label and label not in seen:
            seen.add(label)
            out.append(label)
    return out if cap is None else out[:cap]


def _short(label, n=26):
    return label if len(label) <= n else label[: n - 1] + "\u2026"


# --------------------------------------------------------------------------- #
# 1. Registration timeline heatmap (year x month grid of malicious registrations)
# --------------------------------------------------------------------------- #
_DATE_RX = re.compile(r"(\d{4})[-/](\d{2})(?:[-/](\d{2}))?")


def _parse_ymd(value):
    if not value:
        return None
    m = _DATE_RX.search(str(value))
    if not m:
        return None
    y, mo = int(m.group(1)), int(m.group(2))
    if 1990 <= y <= 2100 and 1 <= mo <= 12:
        return (y, mo)
    return None


def collect_registrations(json_data, case_dir=None):
    """[(domain, (year, month))] for malicious/operator domains. whois sidecar first,
    then timeline 'registered' events, then any subject created/first_seen field."""
    domains = set(_domains(json_data))
    # roles by label to filter out the impersonated brand / victims
    role_by_label = {}
    for s in json_data.get("subjects", []) or []:
        role_by_label[(s.get("label") or "").strip()] = str(s.get("role", "")).lower()

    regs = {}  # domain -> (y, m)

    # (a) whois sidecars — the authoritative creation date per domain
    if case_dir:
        for path in glob.glob(os.path.join(case_dir, "whois", "*.json")):
            try:
                w = json.load(open(path, encoding="utf-8"))
            except Exception:
                continue
            if not isinstance(w, dict):
                continue
            dom = os.path.splitext(os.path.basename(path))[0]
            ym = _parse_ymd(w.get("created") or w.get("createdDate")
                            or w.get("creation_date"))
            if ym and dom in domains:
                regs.setdefault(dom, ym)

    # (b) timeline events mentioning registration
    for ev in json_data.get("timeline", []) or []:
        text = str(ev.get("event") or ev.get("label") or "")
        if not re.search(r"regist|creat|whois", text, re.I):
            continue
        ym = _parse_ymd(ev.get("date"))
        if not ym:
            continue
        for dom in domains:
            if dom in text and dom not in regs:
                regs[dom] = ym

    # (c) subject-level date fields as last resort
    for s in json_data.get("subjects", []) or []:
        if str(s.get("type", "")).lower() != "domain":
            continue
        dom = (s.get("label") or "").strip()
        if dom in regs:
            continue
        ym = _parse_ymd(s.get("whois_created") or s.get("created")
                        or s.get("first_seen") or s.get("registered"))
        if ym:
            regs[dom] = ym

    return [(d, ym) for d, ym in regs.items()
            if role_by_label.get(d, "") not in _NON_MALICIOUS_ROLES]


from cti_report_figures import registration_heatmap_png, cooccurrence_heatmap_png  # noqa: E402


def add_registration_heatmap(doc, regs):
    """Year x month heatmap; cell = count of malicious domains registered that month."""
    from docx.shared import Inches   # lazy: python-docx only where a Document is in hand
    png = registration_heatmap_png(regs)
    if not png:
        return False
    doc.add_picture(BytesIO(png), width=Inches(6.4))
    doc.paragraphs[-1].alignment = 1
    return True


# --------------------------------------------------------------------------- #
# 2. Domain x indicator co-occurrence (correlation) heatmap
# --------------------------------------------------------------------------- #
def _attr_map(json_data, case_dir=None, cap=DOMAIN_CAP):
    """domain -> set(shared-attribute tokens). From report connections, case_graph indicator edges
    and — when the case dir holds WHOIS sidecars — the registrant e-mail / phone join keys, so the
    matrix shows the rung-1 link that binds an estate and not only the page-level artifacts."""
    domains = set(_domains(json_data, cap))
    attrs = {d: set() for d in domains}

    if case_dir:
        for d in domains:
            wp = os.path.join(case_dir, "whois", d + ".json")
            if not os.path.isfile(wp):
                continue
            try:
                w = json.load(open(wp, encoding="utf-8")) or {}
            except Exception:
                continue
            for key, tag in (("registrant_email", "registrant_email"), ("registrant_phone", "registrant_phone")):
                v = (w.get(key) or "").strip().lower()
                if v and "privacy" not in v and "redacted" not in v and not v.startswith("abuse@"):
                    attrs[d].add(f"{tag}:{v}")

    for c in json_data.get("connections", []) or []:
        rel = str(c.get("relationship") or c.get("rel") or "").lower()
        frm = (c.get("from_id") or c.get("source") or "").strip()
        to = (c.get("to_id") or c.get("target") or "").strip()
        if frm not in domains:
            continue
        if rel in _ATTR_RELATIONSHIPS:
            attrs[frm].add(to if ":" in to else f"{rel}:{to}")

    # enrich with case_graph indicator edges (favicon / css / js / ip / cert)
    if case_dir:
        cg = os.path.join(case_dir, "case_graph.json")
        if os.path.isfile(cg):
            try:
                g = json.load(open(cg, encoding="utf-8"))
            except Exception:
                g = None
            for e in (g or {}).get("edges", []) or []:
                frm = (e.get("source") or e.get("from") or "").strip()
                to = (e.get("target") or e.get("to") or "").strip()
                if frm in attrs and to and re.match(
                        r"^(favicon|css_hash|js_bundle|indicator|ip|cert|dom_skeleton|comment):", to):
                    attrs[frm].add(to)
    # §2.5 false-positive control: an attribute the report itself excludes from the IOC set
    # (ioc_exclude) or that the reference ledger marks benign (a saturated favicon, a provider SPF
    # include) must not appear as a "correlation" — the figure would contradict the text.
    excluded = _excluded_tokens(json_data, case_dir)
    if excluded:
        for d in attrs:
            attrs[d] = {t for t in attrs[d] if t.lower() not in excluded
                        and t.partition(":")[2].lower() not in excluded}
    return attrs


def _excluded_tokens(json_data, case_dir=None):
    """Lower-cased values from the report's ioc_exclude list plus every 'benign' verdict in the
    knowledge reference ledger (looked up beside the case dir: <root>/knowledge/reference.jsonl)."""
    out = set()
    for v in json_data.get("ioc_exclude") or []:
        s = v if isinstance(v, str) else (v.get("value") if isinstance(v, dict) else None)
        if s:
            out.add(str(s).lower())
    if case_dir:
        ref = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(case_dir))), "knowledge", "reference.jsonl")
        if os.path.isfile(ref):
            for line in open(ref, encoding="utf-8"):
                try:
                    r = json.loads(line)
                except ValueError:
                    continue
                if r.get("verdict") == "benign" and r.get("value"):
                    out.add(str(r["value"]).lower())
    return out


def build_cooccurrence(json_data, case_dir=None, cap=DOMAIN_CAP):
    """(domains, attr_labels, matrix 0/1). Only attributes shared by >=2 domains."""
    attrs = _attr_map(json_data, case_dir, cap)
    domains = [d for d in _domains(json_data, cap) if attrs.get(d)]
    if len(domains) < 2:
        return None
    counts = {}
    for d in domains:
        for a in attrs[d]:
            counts[a] = counts.get(a, 0) + 1
    shared = [a for a, n in counts.items() if n >= 2]
    if not shared:
        return None
    shared = sorted(shared, key=lambda a: (-counts[a], a))[:20]
    mat = np.zeros((len(domains), len(shared)), dtype=int)
    for i, d in enumerate(domains):
        for j, a in enumerate(shared):
            mat[i, j] = 1 if a in attrs[d] else 0
    return domains, shared, mat


def _attr_label(token):
    kind, _, val = token.partition(":")
    pretty = {"ns": "nameserver", "registrant_email": "registrant email", "registrant_phone": "registrant phone",
              "registrant": "registrant", "registrar": "registrar",
              "css_hash": "CSS", "js_bundle": "JS bundle", "favicon": "favicon",
              "dom_skeleton": "DOM template", "comment": "HTML comment", "ip": "IP"}.get(kind, kind)
    tail = val.split("/")[-1]
    if len(tail) > 14:
        tail = tail[:8] + "\u2026" + tail[-4:]
    return f"{pretty}:{tail}" if tail else pretty


def add_cooccurrence_heatmap(doc, built):
    """domains (rows) x shared indicators (cols); filled cell = domain carries indicator."""
    from docx.shared import Inches
    png = cooccurrence_heatmap_png(built, _attr_label)
    if not png:
        return False
    doc.add_picture(BytesIO(png), width=Inches(6.4))
    doc.paragraphs[-1].alignment = 1
    return True


# --------------------------------------------------------------------------- #
# 3. Domain x domain possible-relation strength heatmap
# --------------------------------------------------------------------------- #
def build_relation_strength(json_data, case_dir=None):
    """(domains, matrix) symmetric; cell = shared indicators + direct sibling/edge links."""
    attrs = _attr_map(json_data, case_dir)
    domains = [d for d in _domains(json_data) if attrs.get(d)]
    if len(domains) < 2:
        # fall back to any domain that participates in a direct edge
        domains = _domains(json_data)
        if len(domains) < 2:
            return None
    idx = {d: i for i, d in enumerate(domains)}
    n = len(domains)
    mat = np.zeros((n, n), dtype=float)

    # shared-attribute weight
    for i in range(n):
        for j in range(i + 1, n):
            shared = len(attrs.get(domains[i], set()) & attrs.get(domains[j], set()))
            mat[i, j] += shared
            mat[j, i] += shared

    # direct edges (sibling / any domain->domain connection)
    for c in json_data.get("connections", []) or []:
        frm = (c.get("from_id") or c.get("source") or "").strip()
        to = (c.get("to_id") or c.get("target") or "").strip()
        if frm in idx and to in idx and frm != to:
            w = 2.0 if str(c.get("relationship") or c.get("rel") or "").lower() == "sibling" else 1.0
            mat[idx[frm], idx[to]] += w
            mat[idx[to], idx[frm]] += w

    if mat.max() <= 0:
        return None
    np.fill_diagonal(mat, np.nan)  # blank diagonal
    return domains, mat


def add_relation_heatmap(doc, built):
    """Symmetric domain x domain relation-strength matrix (higher = more shared evidence)."""
    if not built:
        return False
    domains, mat = built
    n = len(domains)
    fig, ax = plt.subplots(figsize=(max(4.5, 0.42 * n + 2.0),
                                    max(4.0, 0.42 * n + 1.6)), dpi=150)
    vmax = np.nanmax(mat) if np.isfinite(np.nanmax(mat)) else 1
    im = ax.imshow(mat, cmap="YlOrRd", aspect="auto", vmin=0, vmax=max(1, vmax))
    ax.set_xticks(range(n)); ax.set_yticks(range(n))
    ax.set_xticklabels([_short(d, 18) for d in domains], rotation=45, ha="right", fontsize=6.5)
    ax.set_yticklabels([_short(d, 18) for d in domains], fontsize=6.5)
    for i in range(n):
        for j in range(n):
            v = mat[i, j]
            if np.isfinite(v) and v > 0:
                ax.text(j, i, f"{v:.0f}", ha="center", va="center", fontsize=6.5,
                        color="white" if v > max(1, vmax) * 0.6 else COLORS_HEX["text"])
    ax.set_title("Possible-relation strength (shared indicators + direct links)",
                 fontsize=10, color=COLORS_HEX["text"], pad=8)
    cbar = fig.colorbar(im, ax=ax, fraction=0.045, pad=0.02)
    cbar.ax.tick_params(labelsize=7)
    fig.tight_layout()
    from docx.shared import Inches
    from cti_report_figures import fig_png
    doc.add_picture(BytesIO(fig_png(fig)), width=Inches(5.8))
    doc.paragraphs[-1].alignment = 1
    return True
