"""
CTI Report Charts — pie, bar, gauge, timeline via matplotlib → BytesIO → docx.
"""
import matplotlib
matplotlib.use("Agg")

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from io import BytesIO
from docx.shared import Inches

from cti_docx_styles import COLORS_HEX, SEVERITY_COLORS_HEX, CYCLE_HEX


def _save_fig_to_buffer(fig) -> BytesIO:
    """Save matplotlib figure to BytesIO buffer."""
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150, facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf


def add_finding_type_pie(doc, findings: list) -> None:
    """Pie chart: finding distribution by type."""
    if not findings:
        return

    type_counts = {}
    for f in findings:
        t = f.get("type", "unknown")
        type_counts[t] = type_counts.get(t, 0) + 1

    labels = list(type_counts.keys())
    sizes = list(type_counts.values())
    colors = (CYCLE_HEX * ((len(labels) // len(CYCLE_HEX)) + 1))[:len(labels)]

    fig, ax = plt.subplots(figsize=(5, 4), dpi=150)
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, autopct="%1.0f%%",
        colors=colors, startangle=90,
        pctdistance=0.75, labeldistance=1.15,
        textprops={"fontsize": 9}
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontweight("bold")
    ax.set_title("Finding Distribution by Type", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15)

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(4.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1  # CENTER


def add_severity_bar(doc, findings: list) -> None:
    """Horizontal bar chart: finding count by severity."""
    if not findings:
        return

    sev_order = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]
    sev_counts = {s: 0 for s in sev_order}
    for f in findings:
        w = f.get("weight", "INFO").upper()
        if w in sev_counts:
            sev_counts[w] += 1

    labels = [s for s in sev_order if sev_counts[s] > 0]
    counts = [sev_counts[s] for s in labels]
    colors = [SEVERITY_COLORS_HEX.get(s, COLORS_HEX["muted"]) for s in labels]

    fig, ax = plt.subplots(figsize=(5.5, max(2.5, len(labels) * 0.6)), dpi=150)
    bars = ax.barh(labels, counts, color=colors, height=0.5, edgecolor="white")

    for bar, count in zip(bars, counts):
        ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height() / 2,
                str(count), va="center", fontsize=10, fontweight="bold",
                color=COLORS_HEX["text"])

    ax.set_xlabel("Count", fontsize=10, color=COLORS_HEX["muted"])
    ax.set_title("Findings by Severity", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15)
    ax.invert_yaxis()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False)

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


_ICD203 = [
    ("Almost\nno chance", 2), ("Very\nunlikely", 12), ("Unlikely", 32),
    ("Roughly\neven", 50), ("Likely", 67), ("Very\nlikely", 87),
    ("Almost\ncertain", 97),
]
_ADMIRALTY = [
    ("A \u00b7 Completely reliable", ), ("B \u00b7 Usually reliable", ),
    ("C \u00b7 Fairly reliable", ), ("D \u00b7 Not usually reliable", ),
    ("E \u00b7 Unreliable", ), ("F \u00b7 Cannot be judged", ),
]


def _likelihood_col(conf: int) -> int:
    c = max(0, min(100, int(conf)))
    for i, hi in enumerate((5, 20, 45, 55, 80, 95)):
        if c < hi:
            return i
    return 6


def _reliability_row(verified) -> int:
    # Admiralty source reliability derived from corroboration state.
    if verified is True:
        return 1   # B — usually reliable (independently corroborated)
    if verified is False:
        return 3   # D — not usually reliable (single, uncorroborated)
    return 2       # C — fairly reliable (default OSINT posture)


def add_confidence_matrix(doc, findings: list, subjects: list) -> None:
    """Two-axis confidence explainer: ICD-203 likelihood (x) vs Admiralty source
    reliability (y). The labelled grid IS the legend; each finding is plotted in its
    cell, coloured by severity, so a reader sees both the scale and where the case's
    evidence sits on it."""
    ncol, nrow = len(_ICD203), len(_ADMIRALTY)
    fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=150)

    # faint reference grid
    ax.imshow(np.zeros((nrow, ncol)), cmap="Greys", vmin=0, vmax=1, aspect="auto")
    ax.set_xticks(np.arange(-.5, ncol, 1), minor=True)
    ax.set_yticks(np.arange(-.5, nrow, 1), minor=True)
    ax.grid(which="minor", color="#c8cdd3", linewidth=0.8)
    ax.tick_params(which="minor", length=0)

    ax.set_xticks(range(ncol))
    ax.set_xticklabels(["%s\n%d%%" % (lbl, pct) for lbl, pct in _ICD203], fontsize=7)
    ax.set_yticks(range(nrow))
    ax.set_yticklabels([a[0] for a in _ADMIRALTY], fontsize=8)
    ax.set_xlabel("ICD-203 likelihood  \u2014  probability the judgment is correct",
                  fontsize=8.5, color=COLORS_HEX["text"])
    ax.set_ylabel("Admiralty  \u2014  source reliability", fontsize=8.5,
                  color=COLORS_HEX["text"])
    ax.set_title("Confidence scale (two-axis): ICD-203 likelihood \u00d7 Admiralty reliability",
                 fontsize=10, color=COLORS_HEX["text"], pad=10)

    # plot findings, jittered within a cell so co-located markers stay visible
    sub_verified = {s.get("id"): s.get("verified") for s in (subjects or [])}
    rng = np.random.default_rng(7)
    seen_counts = {}
    handles = {}
    for f in findings or []:
        col = _likelihood_col(f.get("confidence", 0))
        row = _reliability_row(sub_verified.get(f.get("subject_id")))
        k = (row, col)
        seen_counts[k] = seen_counts.get(k, 0) + 1
        jx = (rng.random() - 0.5) * 0.5 if seen_counts[k] > 1 else 0.0
        jy = (rng.random() - 0.5) * 0.5 if seen_counts[k] > 1 else 0.0
        sev = str(f.get("weight", "INFO")).upper()
        color = SEVERITY_COLORS_HEX.get(sev, COLORS_HEX["muted"])
        ax.scatter(col + jx, row + jy, s=150, color=color, edgecolor="white",
                   linewidth=1.2, zorder=3)
        handles.setdefault(sev, color)

    if handles:
        order = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]
        patches = [mpatches.Patch(color=handles[s], label=s)
                   for s in order if s in handles]
        ax.legend(handles=patches, title="Finding severity", fontsize=7,
                  title_fontsize=7.5, loc="upper left", bbox_to_anchor=(1.01, 1.0),
                  frameon=False)

    ax.set_xlim(-0.5, ncol - 0.5)
    ax.set_ylim(nrow - 0.5, -0.5)  # A at top
    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(6.6))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


# The Event Timeline lives in its own module (cti_docx_timeline_chart.py); re-exported here so
# cti_docx_postprocess keeps importing add_timeline_chart from this namespace.
from cti_docx_timeline_chart import add_timeline_chart  # noqa: E402,F401

def add_traffic_sources_bar(doc, traffic_sources: dict) -> None:
    """Horizontal bar chart: traffic source breakdown (direct/search/referral/social/paid)."""
    if not traffic_sources:
        return

    source_colors = {
        "direct": "#22333F",    # slate
        "search": "#3B5566",    # steel
        "referral": "#5A6B3B",  # olive
        "social": "#5A4A7A",    # muted violet
        "paid": "#B0790F",      # ochre
        "email": "#7A2F52",     # muted plum
        "display": "#9A5B2F",   # amber-brown
    }

    labels = [k.title() for k in traffic_sources.keys()]
    values = list(traffic_sources.values())
    colors = [source_colors.get(k, COLORS_HEX["muted"]) for k in traffic_sources.keys()]

    fig, ax = plt.subplots(figsize=(5.5, max(2, len(labels) * 0.5)), dpi=150)
    bars = ax.barh(labels, values, color=colors, height=0.5, edgecolor="white")

    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2,
                f"{val}%", va="center", fontsize=10, fontweight="bold",
                color=COLORS_HEX["text"])

    ax.set_xlabel("Percentage (%)", fontsize=10, color=COLORS_HEX["muted"])
    ax.set_title("Traffic Sources", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15)
    ax.set_xlim(0, max(values) * 1.3 if values else 100)
    ax.invert_yaxis()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False)

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


def add_geographic_pie(doc, top_countries: list) -> None:
    """Pie chart: visitor geographic distribution by country."""
    if not top_countries:
        return

    labels = [c.get("country", "?") for c in top_countries]
    sizes = [c.get("share", 0) for c in top_countries]

    # Add "Other" if shares don't sum to 100
    total = sum(sizes)
    if total < 100:
        labels.append("Other")
        sizes.append(100 - total)

    geo_colors = (CYCLE_HEX * ((len(labels) // len(CYCLE_HEX)) + 1))[:len(labels)]

    fig, ax = plt.subplots(figsize=(5, 4), dpi=150)
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, autopct="%1.0f%%",
        colors=geo_colors, startangle=90,
        pctdistance=0.75, labeldistance=1.15,
        textprops={"fontsize": 9}
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontweight("bold")
    ax.set_title("Visitor Geography", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15)

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(4.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1
