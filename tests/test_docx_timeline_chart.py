"""Regression: the DOCX Event Timeline must never truncate labels, must group dated events by
month, must paginate so no single picture exceeds one page, and must not crash on undated or
empty input. (Bug: a 32-event estate rendered as a 13-inch single column with every label cut at
40 chars — "registered through Glob".)"""
import datetime
import os
import random
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, "scripts"))

from docx import Document
import cti_docx_charts  # re-export path used by cti_docx_postprocess
import cti_docx_timeline_chart as charts

_BOILER = " registered through Example Registrar LLC"


def _events(n, start=datetime.date(2025, 1, 1), spread_days=400, seed=1):
    rng = random.Random(seed)
    return [{"date": (start + datetime.timedelta(days=rng.randint(0, spread_days))).isoformat(),
             "event": f"host{i}.example.com{_BOILER}"} for i in range(n)]


def _pictures(doc):
    return [r.target_part.blob for r in doc.part.rels.values() if "image" in r.reltype]


def _png_size(blob):
    # PNG IHDR: width/height are big-endian uint32 at bytes 16..24
    return int.from_bytes(blob[16:20], "big"), int.from_bytes(blob[20:24], "big")


def test_common_suffix_factors_shared_boilerplate_only_when_all_share_it():
    labels = [f"h{i}.example.com{_BOILER}" for i in range(5)]
    assert charts._common_suffix(labels) == _BOILER.strip()
    assert charts._common_suffix(labels + ["h9.example.com moved behind CDN"]) == ""
    assert charts._common_suffix(labels[:2]) == ""  # <3 strings: never factor


def test_wrap_caps_at_two_lines_and_marks_the_cut():
    out = charts._wrap(("word " * 60).strip(), 30)
    assert out.count("\n") + 1 <= charts._TIMELINE_MAX_LABEL_LINES
    assert out.endswith("…")
    assert charts._wrap("vnvc-vn.example", 30) == "vnvc-vn.example"


def test_parse_date_accepts_iso_prefixes_and_rejects_garbage():
    assert charts._parse_date("2025-09-25") == datetime.date(2025, 9, 25)
    assert charts._parse_date("2025-09-25T10:00:00Z") == datetime.date(2025, 9, 25)
    assert charts._parse_date("2025-09") == datetime.date(2025, 9, 1)
    assert charts._parse_date("N/A") is None
    assert charts._parse_date("2026-Q3") is None
    assert charts._parse_date(None) is None


def test_re_export_from_charts_module_is_the_same_function():
    assert cti_docx_charts.add_timeline_chart is charts.add_timeline_chart


def test_single_month_burst_splits_across_pages_with_continuation():
    doc = Document()
    charts.add_timeline_chart(doc, _events(70, start=datetime.date(2025, 5, 1), spread_days=27))
    pics = _pictures(doc)
    assert len(pics) >= 2
    for blob in pics:
        w, h = _png_size(blob)
        assert h / w * 6.3 < 9.2


def test_thirty_two_events_render_within_page_height():
    doc = Document()
    charts.add_timeline_chart(doc, _events(32, spread_days=330))
    pics = _pictures(doc)
    assert pics, "no picture emitted"
    for blob in pics:
        w, h = _png_size(blob)
        # placed at 6.3 in wide; scaled height must fit a Letter page body (< 9.2 in)
        assert h / w * 6.3 < 9.2, f"picture too tall for one page: {h / w * 6.3:.1f} in"


def test_sixty_events_paginate_into_multiple_pictures():
    doc = Document()
    charts.add_timeline_chart(doc, _events(60, spread_days=700))
    pics = _pictures(doc)
    assert len(pics) >= 2
    for blob in pics:
        w, h = _png_size(blob)
        assert h / w * 6.3 < 9.2


def test_undated_and_empty_inputs_do_not_crash():
    doc = Document()
    charts.add_timeline_chart(doc, [])
    assert not _pictures(doc)
    doc = Document()
    charts.add_timeline_chart(doc, _events(4) + [{"date": "N/A", "event": "undated note"},
                                                {"event": "no date key at all"}])
    assert len(_pictures(doc)) == 2  # dated swimlane + undated list
    doc = Document()
    charts.add_timeline_chart(doc, [{"date": "N/A", "event": "only undated"}])
    assert len(_pictures(doc)) == 1  # list fallback


_TESTS = [
    test_re_export_from_charts_module_is_the_same_function,
    test_single_month_burst_splits_across_pages_with_continuation,
    test_common_suffix_factors_shared_boilerplate_only_when_all_share_it,
    test_wrap_caps_at_two_lines_and_marks_the_cut,
    test_parse_date_accepts_iso_prefixes_and_rejects_garbage,
    test_thirty_two_events_render_within_page_height,
    test_sixty_events_paginate_into_multiple_pictures,
    test_undated_and_empty_inputs_do_not_crash,
]


def check():
    passed = failed = 0
    lines = []
    for test in _TESTS:
        label = test.__name__.removeprefix("test_").replace("_", " ")
        try:
            test()
        except Exception as exc:  # noqa: BLE001 — report every independent contract
            failed += 1
            lines.append(("FAIL", f"{label}: {exc}"))
        else:
            passed += 1
            lines.append(("ok", label))
    return passed, failed, lines


if __name__ == "__main__":
    _passed, _failed, _lines = check()
    for _status, _label in _lines:
        print(("ok" if _status == "ok" else "FAIL") + "  " + _label)
    raise SystemExit(bool(_failed))
